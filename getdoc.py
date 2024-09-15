from docx import Document
import requests
from cryptography.hazmat.primitives.ciphers import Cipher, algorithms, modes
import base64
def decrypt_data(encrypted_data: str, key):
    try:
        # Decode the Base64-encoded encrypted data
        encrypted_data_bytes = base64.b64decode(encrypted_data)

        # Decrypt AES-128 ECB encryption (no IV needed)
        # cipher = Cipher(algorithms.AES(key), modes.ECB())
        # decryptor = cipher.decryptor()
        # decrypted_data = decryptor.update(encrypted_data_bytes) + decryptor.finalize()

        return encrypted_data_bytes
    except Exception as e:
        raise Exception(f"Decryption error: {e}")
def read_and_edit_docx(input_file, output_file):
    # Load the .docx file
    doc = Document(input_file)
    all_text=[]
    # Loop through paragraphs and append the text to the list
    for paragraph in doc.paragraphs:
        all_text.append(paragraph.text)


    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            all_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text.append(cell.text)

    for section in doc.sections:
        footer = section.footer
        for paragraph in footer.paragraphs:
            all_text.append(paragraph.text)

    all_text = '\n'.join(all_text)

    try:
        # Prepare payload for API request
        payload = {
            "text": [all_text],
            "link_batch": False,
            "entity_detection": {
                "entity_types": [
                    {
                        "type": "ENABLE",
                        "value": []
                    }
                ]
            },
            "processed_text": {
                "type": "MARKER",
                "pattern": "[UNIQUE_NUMBERED_ENTITY_TYPE]"
            }
        }

        headers = {"Content-Type": "application/json", "X-API-KEY": "a12108691d914f70a90fd99dc595f4af"}

        # Send request to the API
        response = requests.post(os.environ.get('srver'), json=payload, headers=headers)
        response.raise_for_status()
        data = response.json()
        text_to_redact = [(word["text"],word["processed_text"]) for word in data[0]["entities"]]
        dicti={}
        for i,j in text_to_redact:
            dicti[i]=j
        for paragraph in doc.paragraphs:
            for key in dicti.keys():
                paragraph.text = paragraph.text.replace(key,dicti[key])


        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                for key in dicti.keys():
                    paragraph.text = paragraph.text.replace(key,dicti[key])

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                 for key in dicti.keys():
                    cell.text = cell.text.replace(key,dicti[key])

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for key in dicti.keys():
                    paragraph.text = paragraph.text.replace(key,dicti[key])

        # Save the modified document
        doc.save(output_file)

    except requests.RequestException as e:
        logging.error(f"API request failed: {e}")
        raise
    
def run(req):
    encrypted_file = req.docx
    decrypted_file =  decrypt_data(encrypted_file, "1234567890123456")
    with open("output_file.docx", "w", encoding="utf-8") as f:
        f.write(decrypted_file.decode("utf-8"))
# run(req)
# read_and_edit_docx("output_file.docx", "final.docx")
    
    
    